from docx import Document
import json
import os
import re

JSON_FILE = "mapping_data.json"
OUTPUT_FILE = "outputt.docx"

def normalize(text):
    """Lowercase, remove punctuation and extra spaces for matching."""
    return re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', text)).strip().lower()

def load_mapping(json_path):
    """Load JSON as normalized description -> field type -> value -> mnemonic"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    mapping = {}
    for item in data:
        fieldname = item.get("fieldname", "")
        value = str(item.get("value", "")).strip()
        mnemonic = str(item.get("mnemonic", "")).strip()
        parts = fieldname.split(" - ")
        if len(parts) >= 3:
            description = normalize(parts[1].strip())
            field_type = parts[-1].strip()
            mapping.setdefault(description, {}).setdefault(field_type, {})[normalize(value)] = mnemonic
    return mapping

def find_best_match(table_desc, mapping_keys):
    """Find JSON description that is inside table description or vice versa."""
    table_desc_norm = normalize(table_desc)
    for key in mapping_keys:
        if key in table_desc_norm or table_desc_norm in key:
            return key
    return None

def replace_in_table(table, mapping):
    mapping_keys = list(mapping.keys())
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 4:
            table_desc = cells[1].text.strip()
            matched_desc = find_best_match(table_desc, mapping_keys)
            if matched_desc:
                # Amount of Insurance
                current_amount = normalize(cells[2].text.strip())
                if "Amount of Insurance" in mapping[matched_desc]:
                    amt_dict = mapping[matched_desc]["Amount of Insurance"]
                    if current_amount in amt_dict:
                        cells[2].text = amt_dict[current_amount]
                # Premium
                current_premium = normalize(cells[3].text.strip())
                if "Premium" in mapping[matched_desc]:
                    prem_dict = mapping[matched_desc]["Premium"]
                    if current_premium in prem_dict:
                        cells[3].text = prem_dict[current_premium]

def process_doc(doc, mapping):
    # Main body tables
    for table in doc.tables:
        replace_in_table(table, mapping)
    # Headers and footers
    for section in doc.sections:
        for container in [section.header, section.footer]:
            for table in container.tables:
                replace_in_table(table, mapping)

def main():
    input_file = input("Enter DOCX file name: ").strip()
    if not os.path.isfile(input_file):
        print(f"File '{input_file}' not found!")
        return
    if not os.path.isfile(JSON_FILE):
        print(f"Mapping JSON '{JSON_FILE}' not found!")
        return

    mapping = load_mapping(JSON_FILE)
    if not mapping:
        print("Mapping is empty!")
        return

    doc = Document(input_file)
    process_doc(doc, mapping)
    doc.save(OUTPUT_FILE)
    print(f"All replacements done! Output saved as '{OUTPUT_FILE}'")

if __name__ == "__main__":
    main()
